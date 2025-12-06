import os
import sqlite3
from datetime import datetime
from flask import Flask, g, render_template, request, redirect, url_for, flash
from flask import send_file
from docx import Document
import tempfile


# ---------------------------------------------------------------------------
# App factory / config
# ---------------------------------------------------------------------------

def create_app():
    app = Flask(__name__)
    app.config.from_mapping(
        SECRET_KEY="dev",  # change for production
        DATABASE=os.path.join(app.instance_path, "hardware.db"),
    )

    if not os.path.exists(app.instance_path):
        os.makedirs(app.instance_path)

    # Register DB helpers
    app.teardown_appcontext(close_db)

    # Routes
    register_routes(app)

    return app

# ---------------------------------------------------------------------------
# DB helpers
# ---------------------------------------------------------------------------

def get_db():
    if "db" not in g:
        db_path = g._app.config["DATABASE"] if hasattr(g, "_app") else None
        # g._app is not always set; use current_app instead
        from flask import current_app
        db_path = current_app.config["DATABASE"]
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        g.db = conn
    return g.db

def close_db(e=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()

def init_db():
    from flask import current_app
    db = get_db()
    schema_path = os.path.join(current_app.root_path, "schema.sql")
    with current_app.open_resource("schema.sql") as f:
        db.executescript(f.read().decode("utf-8"))
    db.commit()

# ---------------------------------------------------------------------------
# PYYXXX generator (Procedures)
# ---------------------------------------------------------------------------

def generate_new_procedure_id(db):
    """Generate next procedure ID PYY-XXX based on current year."""
    now = datetime.now()
    yy = now.year % 100  # 2025 -> 25
    yy_str = f"{yy:02d}"

    like_pattern = f"P{yy_str}-%"
    cur = db.execute(
        "SELECT proc_id FROM procedures WHERE proc_id LIKE ? ORDER BY proc_id DESC LIMIT 1",
        (like_pattern,),
    )
    row = cur.fetchone()
    if row is None:
        seq = 1
    else:
        last_id = row["proc_id"]  # e.g. 'P25-007'
        last_seq_str = last_id.split("-")[-1]
        try:
            seq = int(last_seq_str) + 1
        except ValueError:
            seq = 1

    seq_str = f"{seq:03d}"
    return f"P{yy_str}-{seq_str}"


# ---------------------------------------------------------------------------
# HYYXXX generator
# ---------------------------------------------------------------------------

def generate_new_hardware_id(db):
    """Generate next HYYXXX ID based on current year and existing records."""
    now = datetime.now()
    yy = now.year % 100  # 2025 -> 25
    yy_str = f"{yy:02d}"

    # Find max XXX for this year
    like_pattern = f"H{yy_str}%"
    cur = db.execute(
        "SELECT hardware_id FROM hardware WHERE hardware_id LIKE ? ORDER BY hardware_id DESC LIMIT 1",
        (like_pattern,),
    )
    row = cur.fetchone()
    if row is None:
        seq = 1
    else:
        last_id = row["hardware_id"]  # e.g. 'H25027'
        last_seq_str = last_id[-3:]
        try:
            seq = int(last_seq_str) + 1
        except ValueError:
            seq = 1

    seq_str = f"{seq:03d}"
    return f"H{yy_str}{seq_str}"

# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

def register_routes(app):
    @app.before_request
    def attach_app_to_g():
        # Small helper so get_db can access config early
        g._app = app

    @app.cli.command("init-db")
    def init_db_command():
        """Initialize the database."""
        init_db()
        print("Initialized the database.")

    @app.route("/")
    def index():
        return redirect(url_for("hardware_list"))

    @app.route("/hardware")
    def hardware_list():
        db = get_db()
        q = request.args.get("q", "").strip()
        category = request.args.get("category", "").strip()
        status = request.args.get("status", "").strip()

        query = "SELECT * FROM hardware WHERE 1=1"
        params = []

        if q:
            query += " AND (hardware_id LIKE ? OR description LIKE ? OR part_number LIKE ?)"
            like = f"%{q}%"
            params.extend([like, like, like])

        if category:
            query += " AND category = ?"
            params.append(category)

        if status:
            query += " AND status = ?"
            params.append(status)

        query += " ORDER BY hardware_id"

        cur = db.execute(query, params)
        items = cur.fetchall()

        # Get distinct categories and statuses for simple filters
        cats = db.execute("SELECT DISTINCT category FROM hardware WHERE category IS NOT NULL ORDER BY category").fetchall()
        stats = db.execute("SELECT DISTINCT status FROM hardware WHERE status IS NOT NULL ORDER BY status").fetchall()

        return render_template(
            "hardware_list.html",
            items=items,
            q=q,
            category=category,
            status=status,
            categories=[c["category"] for c in cats if c["category"]],
            statuses=[s["status"] for s in stats if s["status"]],
        )

    @app.route("/hardware/<int:id>")
    def hardware_detail(id):
        db = get_db()
        cur = db.execute("SELECT * FROM hardware WHERE id = ?", (id,))
        item = cur.fetchone()
        if item is None:
            flash("Hardware not found.", "error")
            return redirect(url_for("hardware_list"))
        return render_template("hardware_detail.html", item=item)

    @app.route("/hardware/new", methods=["GET", "POST"])
    def hardware_new():
        db = get_db()
        if request.method == "POST":
            description = request.form.get("description", "").strip()
            category = request.form.get("category", "").strip()
            part_number = request.form.get("part_number", "").strip()
            serial_number = request.form.get("serial_number", "").strip()
            status = request.form.get("status", "").strip()
            custodian = request.form.get("custodian", "").strip()
            location = request.form.get("location", "").strip()
            traveler_path = request.form.get("traveler_path", "").strip()

            if not description:
                flash("Description is required.", "error")
                return render_template("hardware_form.html", item=None)

            hardware_id = generate_new_hardware_id(db)
            now = datetime.utcnow().isoformat(timespec="seconds")

            db.execute(
                """
                INSERT INTO hardware
                (hardware_id, description, category, part_number, serial_number,
                 status, custodian, location, traveler_path, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    hardware_id,
                    description,
                    category or None,
                    part_number or None,
                    serial_number or None,
                    status or None,
                    custodian or None,
                    location or None,
                    traveler_path or None,
                    now,
                    now,
                ),
            )
            db.commit()

            flash(f"Created hardware {hardware_id}.", "success")
            cur = db.execute("SELECT id FROM hardware WHERE hardware_id = ?", (hardware_id,))
            row = cur.fetchone()
            return redirect(url_for("hardware_detail", id=row["id"]))

        # GET
        return render_template("hardware_form.html", item=None)

    @app.route("/hardware/<int:id>/edit", methods=["GET", "POST"])
    def hardware_edit(id):
        db = get_db()
        cur = db.execute("SELECT * FROM hardware WHERE id = ?", (id,))
        item = cur.fetchone()
        if item is None:
            flash("Hardware not found.", "error")
            return redirect(url_for("hardware_list"))

        if request.method == "POST":
            description = request.form.get("description", "").strip()
            category = request.form.get("category", "").strip()
            part_number = request.form.get("part_number", "").strip()
            serial_number = request.form.get("serial_number", "").strip()
            status = request.form.get("status", "").strip()
            custodian = request.form.get("custodian", "").strip()
            location = request.form.get("location", "").strip()
            traveler_path = request.form.get("traveler_path", "").strip()

            if not description:
                flash("Description is required.", "error")
                return render_template("hardware_form.html", item=item)

            now = datetime.utcnow().isoformat(timespec="seconds")

            db.execute(
                """
                UPDATE hardware
                SET description = ?, category = ?, part_number = ?, serial_number = ?,
                    status = ?, custodian = ?, location = ?, traveler_path = ?, updated_at = ?
                WHERE id = ?
                """,
                (
                    description,
                    category or None,
                    part_number or None,
                    serial_number or None,
                    status or None,
                    custodian or None,
                    location or None,
                    traveler_path or None,
                    now,
                    id,
                ),
            )
            db.commit()
            flash("Hardware updated.", "success")
            return redirect(url_for("hardware_detail", id=id))

        return render_template("hardware_form.html", item=item)

    @app.route("/procedures/<int:id>/edit", methods=["GET", "POST"])
    def procedure_edit(id):
        db = get_db()
        cur = db.execute("SELECT * FROM procedures WHERE id = ?", (id,))
        item = cur.fetchone()
        if item is None:
            flash("Procedure not found.", "error")
            return redirect(url_for("procedure_list"))
    
        if request.method == "POST":
            title = request.form.get("title", "").strip()
            hardware_id = request.form.get("hardware_id", "").strip()
            revision = request.form.get("revision", "").strip()
            purpose = request.form.get("purpose", "").strip()
            hazards = request.form.get("hazards", "").strip()
            prereqs = request.form.get("prereqs", "").strip()
            steps = request.form.get("steps", "").strip()
    
            if not title:
                flash("Title is required.", "error")
                return render_template("procedure_form.html", item=item)
    
            now = datetime.utcnow().isoformat(timespec="seconds")
    
            db.execute(
                """
                UPDATE procedures
                SET title = ?, hardware_id = ?, revision = ?, purpose = ?, hazards = ?,
                    prereqs = ?, steps = ?, updated_at = ?
                WHERE id = ?
                """,
                (
                    title,
                    hardware_id or None,
                    revision or None,
                    purpose or None,
                    hazards or None,
                    prereqs or None,
                    steps or None,
                    now,
                    id,
                ),
            )
            db.commit()
            flash("Procedure updated.", "success")
            return redirect(url_for("procedure_detail", id=id))
    
        return render_template("procedure_form.html", item=item)


    @app.route("/procedures")
    def procedure_list():
        db = get_db()
        q = request.args.get("q", "").strip()

        query = "SELECT * FROM procedures WHERE 1=1"
        params = []

        if q:
            query += " AND (proc_id LIKE ? OR title LIKE ? OR hardware_id LIKE ?)"
            like = f"%{q}%"
            params.extend([like, like, like])

        query += " ORDER BY proc_id"

        cur = db.execute(query, params)
        items = cur.fetchall()

        return render_template(
            "procedure_list.html",
            items=items,
            q=q,
        )

    @app.route("/procedures/<int:id>")
    def procedure_detail(id):
        db = get_db()
        cur = db.execute("SELECT * FROM procedures WHERE id = ?", (id,))
        item = cur.fetchone()
        if item is None:
            flash("Procedure not found.", "error")
            return redirect(url_for("procedure_list"))
        return render_template("procedure_detail.html", item=item)
    
    @app.route("/procedures/new", methods=["GET", "POST"])
    def procedure_new():
        db = get_db()
        if request.method == "POST":
            title = request.form.get("title", "").strip()
            hardware_id = request.form.get("hardware_id", "").strip()
            revision = request.form.get("revision", "").strip() or "A"
            purpose = request.form.get("purpose", "").strip()
            hazards = request.form.get("hazards", "").strip()
            prereqs = request.form.get("prereqs", "").strip()
            steps = request.form.get("steps", "").strip()
    
            if not title:
                flash("Title is required.", "error")
                return render_template("procedure_form.html", item=None)
    
            proc_id = generate_new_procedure_id(db)
            now = datetime.utcnow().isoformat(timespec="seconds")
    
            db.execute(
                """
                INSERT INTO procedures
                (proc_id, title, hardware_id, revision, purpose, hazards, prereqs, steps, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    proc_id,
                    title,
                    hardware_id or None,
                    revision,
                    purpose or None,
                    hazards or None,
                    prereqs or None,
                    steps or None,
                    now,
                    now,
                ),
            )
            db.commit()
    
            flash(f"Created procedure {proc_id}.", "success")
            cur = db.execute("SELECT id FROM procedures WHERE proc_id = ?", (proc_id,))
            row = cur.fetchone()
            return redirect(url_for("procedure_detail", id=row["id"]))
    
        return render_template("procedure_form.html", item=None)


    @app.route("/procedures/<int:id>/sections", methods=["GET", "POST"])
    def procedure_sections(id):
        db = get_db()
        # Ensure procedure exists
        cur = db.execute("SELECT * FROM procedures WHERE id = ?", (id,))
        proc = cur.fetchone()
        if proc is None:
            flash("Procedure not found.", "error")
            return redirect(url_for("procedure_list"))

        if request.method == "POST":
            title = request.form.get("title", "").strip()
            body = request.form.get("body", "").strip()

            if not title:
                flash("Section title is required.", "error")
            else:
                # Find next order_index
                cur = db.execute(
                    "SELECT COALESCE(MAX(order_index), 0) AS max_ord FROM procedure_sections WHERE procedure_id = ?",
                    (id,),
                )
                row = cur.fetchone()
                next_ord = (row["max_ord"] or 0) + 1

                db.execute(
                    """
                    INSERT INTO procedure_sections (procedure_id, order_index, title, body)
                    VALUES (?, ?, ?, ?)
                    """,
                    (id, next_ord, title, body or None),
                )
                db.commit()
                flash("Section added.", "success")

        cur = db.execute(
            "SELECT * FROM procedure_sections WHERE procedure_id = ? ORDER BY order_index",
            (id,),
        )
        sections = cur.fetchall()

        return render_template("procedure_sections.html", proc=proc, sections=sections)

    @app.route("/procedures/<int:id>/export/docx")
    def procedure_export_docx(id):
        db = get_db()
        cur = db.execute("SELECT * FROM procedures WHERE id = ?", (id,))
        proc = cur.fetchone()
        if proc is None:
            flash("Procedure not found.", "error")
            return redirect(url_for("procedure_list"))

        cur = db.execute(
            "SELECT * FROM procedure_sections WHERE procedure_id = ? ORDER BY order_index",
            (id,),
        )
        sections = cur.fetchall()

        doc = Document()
        doc.add_heading(f"{proc['proc_id']} – {proc['title']}", level=1)

        if proc["revision"]:
            doc.add_paragraph(f"Revision: {proc['revision']}")
        if proc["hardware_id"]:
            doc.add_paragraph(f"Hardware ID: {proc['hardware_id']}")
        if proc["purpose"]:
            doc.add_paragraph(f"Purpose: {proc['purpose']}")

        if proc["hazards"]:
            doc.add_paragraph(f"Hazards: {proc['hazards']}")
        if proc["prereqs"]:
            doc.add_paragraph(f"Prereqs: {proc['prereqs']}")

        doc.add_paragraph("")  # spacing

        for s in sections:
            doc.add_heading(s["title"], level=2)
            if s["body"]:
                for line in s["body"].splitlines():
                    doc.add_paragraph(line)

        # If no sections, fall back to steps text
        if not sections and proc["steps"]:
            doc.add_heading("Procedure", level=2)
            for line in proc["steps"].splitlines():
                doc.add_paragraph(line)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.flush()

        filename = f"{proc['proc_id']}.docx"
        return send_file(tmp.name, as_attachment=True, download_name=filename)


app = create_app()

if __name__ == "__main__":
    app.run(debug=True, use_reloader=False)
