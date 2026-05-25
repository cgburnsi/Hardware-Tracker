import io
import os
import uuid
from datetime import datetime
from flask import Blueprint, abort, current_app, flash, redirect, render_template, request, send_file, send_from_directory, url_for
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.db import get_db

bp = Blueprint('ha', __name__, url_prefix='/ha')

SEVERITY_LABELS = {
    1: 'Catastrophic',
    2: 'Critical',
    3: 'Marginal',
    4: 'Negligible',
}

PROBABILITY_LABELS = {
    'A': 'Frequent',
    'B': 'Probable',
    'C': 'Occasional',
    'D': 'Remote',
    'E': 'Improbable',
}

DIST_STATEMENTS = [
    ('A', 'Approved for Public Release; Distribution Unlimited'),
    ('B', 'U.S. Government Agencies Only'),
    ('C', 'U.S. Government Agencies and Their Contractors'),
    ('D', 'Department of Defense and U.S. DoD Contractors Only'),
    ('E', 'DoD Components Only'),
    ('F', 'Further Distribution Only as Directed by Controlling DoD Office'),
    ('X', 'U.S. Government Agencies and Private Individuals or Enterprises '
          'Eligible to Obtain Export-Controlled Technical Data'),
]
_DIST_STATEMENTS_MAP = dict(DIST_STATEMENTS)

# Restrictiveness rank: higher = more restrictive audience.
# B (Gov-only) is technically narrower than C (Gov+contractors) but alphabetical
# ordering is used here as a practical simplification; adjust if needed.
DIST_RANK = {'A': 0, 'B': 1, 'C': 2, 'D': 3, 'E': 4, 'F': 5, 'X': 6}

RAC_RANK = {
    '1A': 'High',     '1B': 'High',     '1C': 'High',     '1D': 'High',     '1E': 'Moderate',
    '2A': 'High',     '2B': 'High',     '2C': 'High',     '2D': 'Moderate', '2E': 'Low',
    '3A': 'High',     '3B': 'Moderate', '3C': 'Moderate', '3D': 'Low',      '3E': 'Minimal',
    '4A': 'Moderate', '4B': 'Low',      '4C': 'Low',      '4D': 'Minimal',  '4E': 'Minimal',
}

CONTROL_TYPES = ['Elimination', 'Substitution', 'Engineering', 'Administrative', 'PPE']


def _build_ha_docx(ha, items_with_recs, references, sig_map,
                   ref_cui_marking, ref_top_dist, ref_top_dist_desc):
    doc = Document()

    for section in doc.sections:
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    marking_parts = []
    if ref_cui_marking:
        marking_parts.append(ref_cui_marking)
    if ref_top_dist:
        marking_parts.append(f'DISTRIBUTION STATEMENT {ref_top_dist}: {ref_top_dist_desc}')
    marking_line = '  |  '.join(marking_parts) if marking_parts else 'UNCONTROLLED'
    is_restricted = bool(ref_cui_marking or (ref_top_dist and ref_top_dist != 'A'))

    def _marking_para():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(marking_line)
        run.bold = True
        run.font.size = Pt(8)
        if is_restricted:
            run.font.color.rgb = RGBColor(0x8B, 0x3A, 0x00)
        return p

    _marking_para()

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'HAZARD ANALYSIS\n{ha["ha_id"]}')
    r.bold = True
    r.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(ha['title'] or '')
    r2.font.size = Pt(12)

    doc.add_paragraph()

    info_tbl = doc.add_table(rows=0, cols=4)
    info_tbl.style = 'Table Grid'

    def _info_row(l1, v1, l2='', v2=''):
        row = info_tbl.add_row()
        for idx, txt in enumerate([l1, str(v1 or '—'), l2, str(v2 or '—')]):
            cell = row.cells[idx]
            cell.text = txt
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(9)
            if idx in (0, 2):
                run.bold = True

    _info_row('Revision', ha['revision'], 'Status', (ha['status'] or '').upper())
    _info_row('Organization', ha['organization'], 'Facility / Operation', ha['facility_operation'])
    _info_row('Classification', ha['preliminary_classification'], 'Created', (ha['created_at'] or '')[:10])

    doc.add_paragraph()

    # Narrative fields
    for heading, field in [('Description', 'description'), ('Scope', 'scope'), ('Assumptions', 'assumptions')]:
        val = ha[field]
        if val:
            h = doc.add_heading(heading, level=2)
            h.runs[0].font.size = Pt(11)
            doc.add_paragraph(val).runs[0].font.size = Pt(10)
            doc.add_paragraph()

    # Reference Documents
    if references:
        h = doc.add_heading('Reference Documents', level=2)
        h.runs[0].font.size = Pt(11)
        ref_tbl = doc.add_table(rows=1, cols=5)
        ref_tbl.style = 'Table Grid'
        for i, label in enumerate(['#', 'Doc Number', 'Title', 'Revision', 'Section Cited']):
            cell = ref_tbl.rows[0].cells[i]
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for idx, ref in enumerate(references, 1):
            row = ref_tbl.add_row()
            for i, val in enumerate([str(idx), ref['doc_number'] or '', ref['title'] or '',
                                      ref['revision'] or '', ref['section_cited'] or '']):
                row.cells[i].text = val
                for r in row.cells[i].paragraphs[0].runs:
                    r.font.size = Pt(9)
        doc.add_paragraph()

    # Hazard Items
    doc.add_heading('Hazard Items', level=1)

    for h_num, item in enumerate(items_with_recs, 1):
        hz   = item['item']
        recs = item['recs']

        h = doc.add_heading(f'H-{h_num:02d}: {hz["hazard_title"]}', level=2)
        h.runs[0].font.size = Pt(11)

        hz_tbl = doc.add_table(rows=0, cols=2)
        hz_tbl.style = 'Table Grid'

        def _hz_row(label, val):
            row = hz_tbl.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(val or '—')
            if row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].bold = True
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            for r in row.cells[1].paragraphs[0].runs:
                r.font.size = Pt(9)

        if hz['hazard_description']:
            _hz_row('Description', hz['hazard_description'])
        _hz_row('Cause', hz['cause'])
        _hz_row('Consequence', hz['consequence'])

        init_rac, init_rank = compute_rac(hz['initial_severity'], hz['initial_probability'])
        final_rac, final_rank = compute_rac(hz['final_severity'], hz['final_probability'])

        if hz['initial_severity']:
            _hz_row('Initial Severity',
                    f"{hz['initial_severity']} — {SEVERITY_LABELS.get(hz['initial_severity'], '')}")
        if hz['initial_probability']:
            _hz_row('Initial Probability',
                    f"{hz['initial_probability']} — {PROBABILITY_LABELS.get(hz['initial_probability'], '')}")
        if init_rac:
            _hz_row('Initial RAC', f"{init_rac} ({init_rank})")
        if hz['final_severity']:
            _hz_row('Final Severity',
                    f"{hz['final_severity']} — {SEVERITY_LABELS.get(hz['final_severity'], '')}")
        if hz['final_probability']:
            _hz_row('Final Probability',
                    f"{hz['final_probability']} — {PROBABILITY_LABELS.get(hz['final_probability'], '')}")
        if final_rac:
            _hz_row('Final RAC', f"{final_rac} ({final_rank})")
        _hz_row('Status', 'Closed' if hz['closed'] else 'Open')

        if recs:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Controls & Recommendations').bold = True
            for rec in recs:
                rp = doc.add_paragraph(style='List Bullet')
                rp.add_run(f"Recommendation: {rec['rec']['text']}").bold = True
                for ctrl in rec['controls']:
                    cp = doc.add_paragraph(style='List Bullet 2')
                    prefix = f"[{ctrl['control_type']}] " if ctrl['control_type'] else ''
                    cp.add_run(f"{prefix}{ctrl['description']}")
                    if ctrl['verification']:
                        cp.add_run(f"  Verification: {ctrl['verification']}").italic = True

        doc.add_paragraph()

    # Signatures
    h = doc.add_heading('Signatures', level=1)
    sig_tbl = doc.add_table(rows=2, cols=3)
    sig_tbl.style = 'Table Grid'
    role_labels = {'prepared_by': 'Prepared By', 'safety_review': 'Safety Review',
                   'risk_acceptance': 'Risk Acceptance'}
    roles = ['prepared_by', 'safety_review', 'risk_acceptance']
    for i, role in enumerate(roles):
        cell = sig_tbl.rows[0].cells[i]
        cell.text = role_labels[role]
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for i, role in enumerate(roles):
        sig = sig_map.get(role, {})
        parts = []
        if sig.get('signer_name'): parts.append(sig['signer_name'])
        if sig.get('signer_org'):  parts.append(sig['signer_org'])
        if sig.get('signed_date'): parts.append(f"Date: {sig['signed_date']}")
        sig_tbl.rows[1].cells[i].text = '\n'.join(parts) if parts else '—'
        for r in sig_tbl.rows[1].cells[i].paragraphs[0].runs:
            r.font.size = Pt(9)

    doc.add_paragraph()
    _marking_para()

    return doc


def refs_locked(ha, sig_map):
    if ha['status'] in ('approved', 'superseded'):
        return True
    return any(s.get('signed_date') for s in sig_map.values())


def compute_rac(severity, probability):
    if severity and probability:
        key = f"{severity}{probability}"
        return key, RAC_RANK.get(key)
    return None, None


def generate_ha_id(db):
    now = datetime.now()
    yy = f"{now.year % 100:02d}"
    row = db.execute(
        "SELECT ha_id FROM hazard_analyses WHERE ha_id LIKE ? ORDER BY ha_id DESC LIMIT 1",
        (f"HA{yy}-%",)
    ).fetchone()
    if row is None:
        seq = 1
    else:
        try:
            seq = int(row['ha_id'].split('-')[1]) + 1
        except Exception:
            seq = 1
    return f"HA{yy}-{seq:03d}"


def get_ha_or_404(db, ha_pk):
    ha = db.execute("SELECT * FROM hazard_analyses WHERE id = ?", (ha_pk,)).fetchone()
    if ha is None:
        abort(404)
    return ha


def get_hazard_or_404(db, item_id, ha_pk):
    item = db.execute(
        "SELECT * FROM hazard_items WHERE id = ? AND ha_id = ?", (item_id, ha_pk)
    ).fetchone()
    if item is None:
        abort(404)
    return item


def get_hazard_position(db, item_id, ha_pk):
    rows = db.execute(
        "SELECT id FROM hazard_items WHERE ha_id = ? ORDER BY order_index, id", (ha_pk,)
    ).fetchall()
    for i, row in enumerate(rows, 1):
        if row['id'] == item_id:
            return i
    return 0


def augment_items(items):
    result = []
    for item in items:
        init_rac, init_rank = compute_rac(item['initial_severity'], item['initial_probability'])
        final_rac, final_rank = compute_rac(item['final_severity'], item['final_probability'])
        result.append({
            **dict(item),
            'init_rac': init_rac, 'init_rank': init_rank,
            'final_rac': final_rac, 'final_rank': final_rank,
        })
    return result


# ── LIST ──────────────────────────────────────────────────────────────────────

@bp.route('/')
def ha_list():
    db = get_db()
    analyses = db.execute("""
        SELECT ha.*, COUNT(hi.id) as hazard_count
        FROM hazard_analyses ha
        LEFT JOIN hazard_items hi ON hi.ha_id = ha.id
        GROUP BY ha.id
        ORDER BY ha.created_at DESC
    """).fetchall()
    return render_template('ha_list.html', analyses=analyses)


# ── NEW ───────────────────────────────────────────────────────────────────────

@bp.route('/new', methods=['GET', 'POST'])
def ha_new():
    db = get_db()
    procedures = db.execute(
        "SELECT id, proc_id, title FROM procedures ORDER BY proc_id"
    ).fetchall()
    if request.method == 'POST':
        ha_id = generate_ha_id(db)
        now = datetime.now().isoformat(timespec='seconds')
        db.execute("""
            INSERT INTO hazard_analyses
              (ha_id, title, facility_operation, organization, preliminary_classification,
               description, scope, assumptions, linked_procedure_id, revision, status, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            ha_id,
            request.form['title'].strip(),
            request.form.get('facility_operation', '').strip(),
            request.form.get('organization', '').strip(),
            request.form.get('preliminary_classification', '').strip(),
            request.form.get('description', '').strip(),
            request.form.get('scope', '').strip(),
            request.form.get('assumptions', '').strip(),
            request.form.get('linked_procedure_id') or None,
            request.form.get('revision', 'A').strip() or 'A',
            request.form.get('status', 'draft'),
            now, now,
        ))
        db.commit()
        row = db.execute("SELECT id FROM hazard_analyses WHERE ha_id = ?", (ha_id,)).fetchone()
        _log_ha(db, row['id'], ha_id, 'Created', f"Created: {request.form['title'].strip()}")
        db.commit()
        flash(f"{ha_id} created.", 'success')
        return redirect(url_for('ha.ha_detail', ha_pk=row['id']))
    return render_template('ha_form.html', ha=None, procedures=procedures)


# ── DETAIL ────────────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>')
def ha_detail(ha_pk):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    raw_items = db.execute(
        "SELECT * FROM hazard_items WHERE ha_id = ? ORDER BY order_index, id", (ha_pk,)
    ).fetchall()
    items = augment_items(raw_items)
    signatures = db.execute(
        "SELECT * FROM hazard_signatures WHERE ha_id = ?", (ha_pk,)
    ).fetchall()
    sig_map = {s['role']: dict(s) for s in signatures}
    references = db.execute(
        """SELECT hr.id, hr.ha_id, hr.ref_doc_id, hr.doc_number, hr.title,
                  hr.revision, hr.section_cited, hr.sort_order,
                  rd.cui, rd.cui_categories, rd.cui_dissem, rd.dist_statement, rd.dist_reason
           FROM hazard_references hr
           LEFT JOIN reference_documents rd ON hr.ref_doc_id = rd.id
           WHERE hr.ha_id = ? ORDER BY hr.sort_order, hr.id""",
        (ha_pk,)
    ).fetchall()
    ref_library = db.execute(
        "SELECT * FROM reference_documents ORDER BY sort_order, doc_number"
    ).fetchall()
    locked = refs_locked(ha, sig_map)

    # Collect all markings from linked library docs for the banner
    marking_rows = db.execute(
        """SELECT rd.cui, rd.cui_categories, rd.cui_dissem, rd.dist_statement
           FROM hazard_references hr
           JOIN reference_documents rd ON hr.ref_doc_id = rd.id
           WHERE hr.ha_id = ?""",
        (ha_pk,)
    ).fetchall()

    # CUI roll-up: union of all categories and dissemination controls
    ref_has_cui = any(r['cui'] for r in marking_rows)
    _seen_cats, _seen_dissems = set(), set()
    ref_cui_cats, ref_cui_dissems = [], []
    for r in marking_rows:
        if r['cui']:
            for c in sorted((r['cui_categories'] or '').split(',')):
                c = c.strip()
                if c and c not in _seen_cats:
                    _seen_cats.add(c)
                    ref_cui_cats.append(c)
            for d in sorted((r['cui_dissem'] or '').split(',')):
                d = d.strip()
                if d and d not in _seen_dissems:
                    _seen_dissems.add(d)
                    ref_cui_dissems.append(d)
    ref_cui_marking = None
    if ref_has_cui:
        parts = ['CUI']
        if ref_cui_cats:    parts.append('//' + '/'.join(ref_cui_cats))
        if ref_cui_dissems: parts.append('//' + '/'.join(ref_cui_dissems))
        ref_cui_marking = ''.join(parts)

    # Distribution: pick single most-restrictive statement by DIST_RANK
    dist_codes = {r['dist_statement'] for r in marking_rows if r['dist_statement']}
    ref_top_dist = max(dist_codes, key=lambda c: DIST_RANK.get(c, -1)) if dist_codes else None
    ref_top_dist_desc = _DIST_STATEMENTS_MAP.get(ref_top_dist, '') if ref_top_dist else ''
    ref_has_restricted = ref_has_cui or (ref_top_dist is not None and ref_top_dist != 'A')

    linked_proc = None
    if ha['linked_procedure_id']:
        linked_proc = db.execute(
            "SELECT proc_id, title FROM procedures WHERE id = ?",
            (ha['linked_procedure_id'],)
        ).fetchone()
    cui_categories = db.execute(
        "SELECT code, label, description FROM cui_categories ORDER BY sort_order, code"
    ).fetchall()
    dissem_controls = db.execute(
        "SELECT code, label, description FROM cui_dissem_controls ORDER BY sort_order, code"
    ).fetchall()

    event_log, log_total, log_page, log_per_page, log_total_pages, log_page_list = _paginate_log(
        db, 'hazard_analyses_log',
        ['action_type', 'description'],
        request.args.get('log_q', '').strip(),
        request.args.get('log_per_page', '25'),
        request.args.get('log_page', 1),
        extra_where=f'ha_id = {ha_pk}',
    )
    log_q = request.args.get('log_q', '').strip()

    return render_template('ha_detail.html',
        ha=ha, items=items, sig_map=sig_map, linked_proc=linked_proc,
        references=references, ref_library=ref_library, locked=locked,
        ref_has_cui=ref_has_cui, ref_cui_marking=ref_cui_marking,
        ref_top_dist=ref_top_dist, ref_top_dist_desc=ref_top_dist_desc,
        ref_has_restricted=ref_has_restricted,
        cui_categories=cui_categories, dissem_controls=dissem_controls,
        dist_statements=DIST_STATEMENTS,
        SEVERITY_LABELS=SEVERITY_LABELS, PROBABILITY_LABELS=PROBABILITY_LABELS,
        event_log=event_log, log_q=log_q, log_per_page=log_per_page,
        log_page=log_page, log_total=log_total,
        log_total_pages=log_total_pages, log_page_list=log_page_list,
    )


# ── EDIT ──────────────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/edit', methods=['GET', 'POST'])
def ha_edit(ha_pk):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    procedures = db.execute(
        "SELECT id, proc_id, title FROM procedures ORDER BY proc_id"
    ).fetchall()
    if request.method == 'POST':
        now = datetime.now().isoformat(timespec='seconds')
        db.execute("""
            UPDATE hazard_analyses SET
              title=?, facility_operation=?, organization=?, preliminary_classification=?,
              description=?, scope=?, assumptions=?, linked_procedure_id=?,
              revision=?, status=?, updated_at=?
            WHERE id=?
        """, (
            request.form['title'].strip(),
            request.form.get('facility_operation', '').strip(),
            request.form.get('organization', '').strip(),
            request.form.get('preliminary_classification', '').strip(),
            request.form.get('description', '').strip(),
            request.form.get('scope', '').strip(),
            request.form.get('assumptions', '').strip(),
            request.form.get('linked_procedure_id') or None,
            request.form.get('revision', 'A').strip() or 'A',
            request.form.get('status', 'draft'),
            now, ha_pk,
        ))
        _log_ha(db, ha_pk, ha['ha_id'], 'Edited',
                f"Edited: rev {request.form.get('revision','').strip() or 'A'}, status {request.form.get('status','draft')}")
        db.commit()
        flash('Hazard Analysis updated.', 'success')
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk))
    return render_template('ha_form.html', ha=ha, procedures=procedures)


# ── DELETE ────────────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/delete', methods=['POST'])
def ha_delete(ha_pk):
    db = get_db()
    get_ha_or_404(db, ha_pk)
    item_rows = db.execute("SELECT id FROM hazard_items WHERE ha_id = ?", (ha_pk,)).fetchall()
    for row in item_rows:
        db.execute("DELETE FROM hazard_controls WHERE hazard_item_id = ?", (row['id'],))
        db.execute("DELETE FROM hazard_notes WHERE hazard_item_id = ?", (row['id'],))
    db.execute("DELETE FROM hazard_items WHERE ha_id = ?", (ha_pk,))
    db.execute("DELETE FROM hazard_signatures WHERE ha_id = ?", (ha_pk,))
    db.execute("DELETE FROM hazard_references WHERE ha_id = ?", (ha_pk,))
    db.execute("DELETE FROM hazard_analyses WHERE id = ?", (ha_pk,))
    db.commit()
    flash('Hazard Analysis deleted.', 'success')
    return redirect(url_for('ha.ha_list'))


# ── SIGNATURES ────────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/signatures', methods=['POST'])
def ha_signatures(ha_pk):
    db = get_db()
    get_ha_or_404(db, ha_pk)
    ha = get_ha_or_404(db, ha_pk)
    sig_parts = []
    for role in ('prepared_by', 'safety_review', 'risk_acceptance'):
        name = request.form.get(f'{role}_name', '').strip()
        org  = request.form.get(f'{role}_org',  '').strip()
        date = request.form.get(f'{role}_date', '').strip()
        existing = db.execute(
            "SELECT id FROM hazard_signatures WHERE ha_id = ? AND role = ?", (ha_pk, role)
        ).fetchone()
        if existing:
            db.execute(
                "UPDATE hazard_signatures SET signer_name=?, signer_org=?, signed_date=? WHERE id=?",
                (name, org, date, existing['id'])
            )
        else:
            db.execute(
                "INSERT INTO hazard_signatures (ha_id, role, signer_name, signer_org, signed_date) VALUES (?,?,?,?,?)",
                (ha_pk, role, name, org, date)
            )
        if name:
            sig_parts.append(f"{role.replace('_',' ').title()}: {name}")
    _log_ha(db, ha_pk, ha['ha_id'], 'Signatures',
            '; '.join(sig_parts) if sig_parts else 'Signatures cleared')
    db.commit()
    flash('Signatures saved.', 'success')
    return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#signatures')


# ── HAZARD NEW ────────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/hazards/new', methods=['GET', 'POST'])
def hazard_new(ha_pk):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    if request.method == 'POST':
        row = db.execute(
            "SELECT MAX(order_index) as mx FROM hazard_items WHERE ha_id = ?", (ha_pk,)
        ).fetchone()
        next_idx = (row['mx'] or 0) + 1
        db.execute("""
            INSERT INTO hazard_items
              (ha_id, order_index, hazard_title, hazard_description, cause, consequence,
               initial_severity, initial_probability, final_severity, final_probability, closed)
            VALUES (?,?,?,?,?,?,?,?,?,?,0)
        """, (
            ha_pk, next_idx,
            request.form['hazard_title'].strip(),
            request.form.get('hazard_description', '').strip(),
            request.form.get('cause', '').strip(),
            request.form.get('consequence', '').strip(),
            request.form.get('initial_severity') or None,
            request.form.get('initial_probability') or None,
            request.form.get('final_severity') or None,
            request.form.get('final_probability') or None,
        ))
        db.commit()
        item_id = db.execute("SELECT last_insert_rowid() as id").fetchone()['id']
        title = request.form['hazard_title'].strip()
        _log_ha(db, ha_pk, ha['ha_id'], 'Hazard Added', f"Added hazard: {title}")
        db.commit()
        flash('Hazard added.', 'success')
        return redirect(url_for('ha.hazard_detail', ha_pk=ha_pk, item_id=item_id))
    return render_template('ha_hazard_form.html', ha=ha, item=None,
                           SEVERITY_LABELS=SEVERITY_LABELS, PROBABILITY_LABELS=PROBABILITY_LABELS)


# ── HAZARD DETAIL ─────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/hazards/<int:item_id>')
def hazard_detail(ha_pk, item_id):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    item = get_hazard_or_404(db, item_id, ha_pk)
    position = get_hazard_position(db, item_id, ha_pk)
    raw_recs = db.execute(
        "SELECT * FROM hazard_recommendations WHERE hazard_item_id = ? ORDER BY order_index, id",
        (item_id,)
    ).fetchall()
    recommendations = []
    for rec in raw_recs:
        ctrls = db.execute(
            "SELECT * FROM hazard_controls WHERE recommendation_id = ? ORDER BY order_index, id",
            (rec['id'],)
        ).fetchall()
        recommendations.append({'rec': rec, 'controls': ctrls})
    notes = db.execute(
        "SELECT * FROM hazard_notes WHERE hazard_item_id = ? ORDER BY created_at",
        (item_id,)
    ).fetchall()
    # Adjacent hazard navigation
    all_ids = [r['id'] for r in db.execute(
        "SELECT id FROM hazard_items WHERE ha_id = ? ORDER BY order_index, id", (ha_pk,)
    ).fetchall()]
    pos_idx = all_ids.index(item_id) if item_id in all_ids else -1
    prev_id = all_ids[pos_idx - 1] if pos_idx > 0 else None
    next_id = all_ids[pos_idx + 1] if pos_idx >= 0 and pos_idx < len(all_ids) - 1 else None

    init_rac, init_rank = compute_rac(item['initial_severity'], item['initial_probability'])
    final_rac, final_rank = compute_rac(item['final_severity'], item['final_probability'])
    return render_template('ha_hazard_detail.html',
        ha=ha, item=item, position=position,
        recommendations=recommendations, notes=notes,
        prev_id=prev_id, next_id=next_id,
        init_rac=init_rac, init_rank=init_rank,
        final_rac=final_rac, final_rank=final_rank,
        SEVERITY_LABELS=SEVERITY_LABELS, PROBABILITY_LABELS=PROBABILITY_LABELS,
        CONTROL_TYPES=CONTROL_TYPES,
    )


# ── HAZARD EDIT ───────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/hazards/<int:item_id>/edit', methods=['GET', 'POST'])
def hazard_edit(ha_pk, item_id):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    item = get_hazard_or_404(db, item_id, ha_pk)
    if request.method == 'POST':
        db.execute("""
            UPDATE hazard_items SET
              hazard_title=?, hazard_description=?, cause=?, consequence=?,
              initial_severity=?, initial_probability=?,
              final_severity=?, final_probability=?
            WHERE id=?
        """, (
            request.form['hazard_title'].strip(),
            request.form.get('hazard_description', '').strip(),
            request.form.get('cause', '').strip(),
            request.form.get('consequence', '').strip(),
            request.form.get('initial_severity') or None,
            request.form.get('initial_probability') or None,
            request.form.get('final_severity') or None,
            request.form.get('final_probability') or None,
            item_id,
        ))
        _log_ha(db, ha_pk, ha['ha_id'], 'Hazard Edited',
                f"Edited hazard: {request.form['hazard_title'].strip()}")
        db.commit()
        flash('Hazard updated.', 'success')
        return redirect(url_for('ha.hazard_detail', ha_pk=ha_pk, item_id=item_id))
    return render_template('ha_hazard_form.html', ha=ha, item=item,
                           SEVERITY_LABELS=SEVERITY_LABELS, PROBABILITY_LABELS=PROBABILITY_LABELS)


# ── HAZARD DELETE ─────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/hazards/<int:item_id>/delete', methods=['POST'])
def hazard_delete(ha_pk, item_id):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    item = get_hazard_or_404(db, item_id, ha_pk)
    rec_rows = db.execute("SELECT id FROM hazard_recommendations WHERE hazard_item_id = ?", (item_id,)).fetchall()
    for rec in rec_rows:
        db.execute("DELETE FROM hazard_controls WHERE recommendation_id = ?", (rec['id'],))
    db.execute("DELETE FROM hazard_recommendations WHERE hazard_item_id = ?", (item_id,))
    db.execute("DELETE FROM hazard_controls WHERE hazard_item_id = ?", (item_id,))
    db.execute("DELETE FROM hazard_notes WHERE hazard_item_id = ?", (item_id,))
    db.execute("DELETE FROM hazard_items WHERE id = ?", (item_id,))
    _log_ha(db, ha_pk, ha['ha_id'], 'Hazard Removed',
            f"Removed hazard: {item['hazard_title'] or '(untitled)'}")
    db.commit()
    flash('Hazard removed.', 'success')
    return redirect(url_for('ha.ha_detail', ha_pk=ha_pk))


# ── HAZARD CLOSE TOGGLE ───────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/hazards/<int:item_id>/close', methods=['POST'])
def hazard_close(ha_pk, item_id):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    item = get_hazard_or_404(db, item_id, ha_pk)
    now_closed = 0 if item['closed'] else 1
    db.execute("UPDATE hazard_items SET closed=? WHERE id=?", (now_closed, item_id))
    action = 'Hazard Closed' if now_closed else 'Hazard Reopened'
    _log_ha(db, ha_pk, ha['ha_id'], action,
            f"{'Closed' if now_closed else 'Reopened'}: {item['hazard_title'] or '(untitled)'}")
    db.commit()
    return redirect(url_for('ha.hazard_detail', ha_pk=ha_pk, item_id=item_id))


# ── NOTES ─────────────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/hazards/<int:item_id>/notes', methods=['POST'])
def note_add(ha_pk, item_id):
    db = get_db()
    get_ha_or_404(db, ha_pk)
    get_hazard_or_404(db, item_id, ha_pk)
    author = request.form.get('author', '').strip()
    body   = request.form.get('body', '').strip()
    if author and body:
        ha = get_ha_or_404(db, ha_pk)
        item = get_hazard_or_404(db, item_id, ha_pk)
        db.execute(
            "INSERT INTO hazard_notes (hazard_item_id, author, body, created_at) VALUES (?,?,?,?)",
            (item_id, author, body, datetime.now().isoformat(timespec='seconds'))
        )
        _log_ha(db, ha_pk, ha['ha_id'], 'Note Added',
                f"Note on '{item['hazard_title'] or '(untitled)'}' by {author}")
        db.commit()
    return redirect(url_for('ha.hazard_detail', ha_pk=ha_pk, item_id=item_id) + '#notes')


# ── RECOMMENDATIONS ───────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/hazards/<int:item_id>/recs/add', methods=['POST'])
def rec_add(ha_pk, item_id):
    db = get_db()
    get_ha_or_404(db, ha_pk)
    get_hazard_or_404(db, item_id, ha_pk)
    text = request.form.get('text', '').strip()
    if text:
        ha = get_ha_or_404(db, ha_pk)
        item = get_hazard_or_404(db, item_id, ha_pk)
        row = db.execute(
            "SELECT MAX(order_index) as mx FROM hazard_recommendations WHERE hazard_item_id=?", (item_id,)
        ).fetchone()
        next_idx = (row['mx'] or 0) + 1
        db.execute(
            "INSERT INTO hazard_recommendations (hazard_item_id, order_index, text) VALUES (?,?,?)",
            (item_id, next_idx, text)
        )
        _log_ha(db, ha_pk, ha['ha_id'], 'Recommendation Added',
                f"Rec on '{item['hazard_title'] or '(untitled)'}': {text[:80]}")
        db.commit()
    return redirect(url_for('ha.hazard_detail', ha_pk=ha_pk, item_id=item_id) + '#controls')


@bp.route('/<int:ha_pk>/hazards/<int:item_id>/recs/<int:rec_id>/delete', methods=['POST'])
def rec_delete(ha_pk, item_id, rec_id):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    item = get_hazard_or_404(db, item_id, ha_pk)
    rec = db.execute("SELECT text FROM hazard_recommendations WHERE id=?", (rec_id,)).fetchone()
    db.execute("DELETE FROM hazard_controls WHERE recommendation_id = ?", (rec_id,))
    db.execute("DELETE FROM hazard_recommendations WHERE id=? AND hazard_item_id=?", (rec_id, item_id))
    _log_ha(db, ha_pk, ha['ha_id'], 'Recommendation Removed',
            f"Removed rec from '{item['hazard_title'] or '(untitled)'}'"
            + (f": {rec['text'][:80]}" if rec else ''))
    db.commit()
    return redirect(url_for('ha.hazard_detail', ha_pk=ha_pk, item_id=item_id) + '#controls')


# ── CONTROLS ──────────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/hazards/<int:item_id>/recs/<int:rec_id>/controls/add', methods=['POST'])
def control_add(ha_pk, item_id, rec_id):
    db = get_db()
    get_ha_or_404(db, ha_pk)
    get_hazard_or_404(db, item_id, ha_pk)
    rec = db.execute(
        "SELECT id FROM hazard_recommendations WHERE id=? AND hazard_item_id=?", (rec_id, item_id)
    ).fetchone()
    if not rec:
        abort(404)
    row = db.execute(
        "SELECT MAX(order_index) as mx FROM hazard_controls WHERE recommendation_id=?", (rec_id,)
    ).fetchone()
    next_idx = (row['mx'] or 0) + 1
    desc = request.form.get('description', '').strip()
    if desc:
        ha = get_ha_or_404(db, ha_pk)
        item = get_hazard_or_404(db, item_id, ha_pk)
        ctrl_type = request.form.get('control_type', '').strip()
        db.execute("""
            INSERT INTO hazard_controls
              (hazard_item_id, recommendation_id, order_index, control_type, description, verification)
            VALUES (?,?,?,?,?,?)
        """, (
            item_id, rec_id, next_idx,
            ctrl_type, desc,
            request.form.get('verification', '').strip(),
        ))
        type_str = f"[{ctrl_type}] " if ctrl_type else ''
        _log_ha(db, ha_pk, ha['ha_id'], 'Control Added',
                f"Control on '{item['hazard_title'] or '(untitled)'}': {type_str}{desc[:80]}")
        db.commit()
    return redirect(url_for('ha.hazard_detail', ha_pk=ha_pk, item_id=item_id) + '#controls')


@bp.route('/<int:ha_pk>/hazards/<int:item_id>/recs/<int:rec_id>/controls/<int:ctrl_id>/delete', methods=['POST'])
def control_delete(ha_pk, item_id, rec_id, ctrl_id):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    item = get_hazard_or_404(db, item_id, ha_pk)
    ctrl = db.execute("SELECT control_type, description FROM hazard_controls WHERE id=?", (ctrl_id,)).fetchone()
    db.execute("DELETE FROM hazard_controls WHERE id=? AND recommendation_id=?", (ctrl_id, rec_id))
    if ctrl:
        type_str = f"[{ctrl['control_type']}] " if ctrl['control_type'] else ''
        _log_ha(db, ha_pk, ha['ha_id'], 'Control Removed',
                f"Removed control from '{item['hazard_title'] or '(untitled)'}': {type_str}{ctrl['description'][:80]}")
    db.commit()
    return redirect(url_for('ha.hazard_detail', ha_pk=ha_pk, item_id=item_id) + '#controls')


# ── EXPORT ────────────────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/export/docx')
def ha_export_docx(ha_pk):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)

    raw_items = db.execute(
        "SELECT * FROM hazard_items WHERE ha_id = ? ORDER BY order_index, id", (ha_pk,)
    ).fetchall()
    items_with_recs = []
    for item in raw_items:
        raw_recs = db.execute(
            "SELECT * FROM hazard_recommendations WHERE hazard_item_id = ? ORDER BY order_index, id",
            (item['id'],)
        ).fetchall()
        recs = []
        for rec in raw_recs:
            ctrls = db.execute(
                "SELECT * FROM hazard_controls WHERE recommendation_id = ? ORDER BY order_index, id",
                (rec['id'],)
            ).fetchall()
            recs.append({'rec': rec, 'controls': ctrls})
        items_with_recs.append({'item': item, 'recs': recs})

    references = db.execute(
        """SELECT hr.doc_number, hr.title, hr.revision, hr.section_cited,
                  rd.cui, rd.cui_categories, rd.cui_dissem, rd.dist_statement
           FROM hazard_references hr
           LEFT JOIN reference_documents rd ON hr.ref_doc_id = rd.id
           WHERE hr.ha_id = ? ORDER BY hr.sort_order, hr.id""",
        (ha_pk,)
    ).fetchall()

    signatures = db.execute("SELECT * FROM hazard_signatures WHERE ha_id = ?", (ha_pk,)).fetchall()
    sig_map = {s['role']: dict(s) for s in signatures}

    marking_rows = db.execute(
        """SELECT rd.cui, rd.cui_categories, rd.cui_dissem, rd.dist_statement
           FROM hazard_references hr
           JOIN reference_documents rd ON hr.ref_doc_id = rd.id
           WHERE hr.ha_id = ?""",
        (ha_pk,)
    ).fetchall()

    ref_has_cui = any(r['cui'] for r in marking_rows)
    _seen_cats, _seen_dissems = set(), set()
    ref_cui_cats, ref_cui_dissems = [], []
    for r in marking_rows:
        if r['cui']:
            for c in sorted((r['cui_categories'] or '').split(',')):
                c = c.strip()
                if c and c not in _seen_cats:
                    _seen_cats.add(c)
                    ref_cui_cats.append(c)
            for d in sorted((r['cui_dissem'] or '').split(',')):
                d = d.strip()
                if d and d not in _seen_dissems:
                    _seen_dissems.add(d)
                    ref_cui_dissems.append(d)
    ref_cui_marking = None
    if ref_has_cui:
        parts = ['CUI']
        if ref_cui_cats:    parts.append('//' + '/'.join(ref_cui_cats))
        if ref_cui_dissems: parts.append('//' + '/'.join(ref_cui_dissems))
        ref_cui_marking = ''.join(parts)

    dist_codes = {r['dist_statement'] for r in marking_rows if r['dist_statement']}
    ref_top_dist = max(dist_codes, key=lambda c: DIST_RANK.get(c, -1)) if dist_codes else None
    ref_top_dist_desc = _DIST_STATEMENTS_MAP.get(ref_top_dist, '') if ref_top_dist else ''

    docx_doc = _build_ha_docx(
        ha, items_with_recs, references, sig_map,
        ref_cui_marking, ref_top_dist, ref_top_dist_desc,
    )

    buf = io.BytesIO()
    docx_doc.save(buf)
    buf.seek(0)
    filename = f"{ha['ha_id']}_Rev{ha['revision'] or 'A'}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# ── REFERENCE LIBRARY (config) ────────────────────────────────────────────────

_ALLOWED_REF_DOC_EXT = {
    'pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx',
    'txt', 'csv', 'dwg', 'dxf', 'zip', 'png', 'jpg', 'jpeg',
}

def _parse_cui_categories(form):
    checked = form.getlist('cui_categories')
    custom  = [c.strip().upper() for c in form.get('cui_custom', '').split(',') if c.strip()]
    seen, result = set(), []
    for code in checked + custom:
        if code not in seen:
            seen.add(code)
            result.append(code)
    return ','.join(result) or None


def _build_cui_marking(cui_categories, cui_dissem):
    """Build the full CUI marking string per CUI Registry format."""
    parts = ['CUI']
    cats = [c for c in (cui_categories or '').split(',') if c]
    dissems = [d for d in (cui_dissem or '').split(',') if d]
    if cats:
        parts.append('//' + '/'.join(cats))
    if dissems:
        parts.append('//' + '/'.join(dissems))
    return ''.join(parts)


def _parse_cui_dissem(form):
    checked = form.getlist('cui_dissem')
    custom  = [c.strip().upper() for c in form.get('cui_dissem_custom', '').split(',') if c.strip()]
    seen, result = set(), []
    for code in checked + custom:
        if code not in seen:
            seen.add(code)
            result.append(code)
    return ','.join(result) or None


def _log_ha(db, ha_pk, ha_code, action_type, description):
    db.execute(
        "INSERT INTO hazard_analyses_log (ha_id, ha_code, timestamp, action_type, description)"
        " VALUES (?, ?, ?, ?, ?)",
        (ha_pk, ha_code, datetime.utcnow().isoformat(timespec='seconds'), action_type, description)
    )


def _log_ref_doc(db, ref_doc_id, doc_number, action_type, description):
    db.execute(
        "INSERT INTO ref_doc_log (ref_doc_id, doc_number, timestamp, action_type, description)"
        " VALUES (?, ?, ?, ?, ?)",
        (ref_doc_id, doc_number, datetime.utcnow().isoformat(timespec='seconds'), action_type, description)
    )


def _ref_doc_folder():
    folder = os.path.join(current_app.instance_path, 'uploads', 'ref_docs')
    os.makedirs(folder, exist_ok=True)
    return folder


@bp.route('/ref-docs/file/<path:stored_name>')
def serve_ref_doc_file(stored_name):
    return send_from_directory(_ref_doc_folder(), stored_name)


_REF_DOC_SORT_COLS  = {'doc_number', 'title', 'revision', 'effective_date', 'expiration_date'}
_VALID_REF_PER_PAGE = (10, 25, 50, 100)
_VALID_LOG_PER_PAGE = {10, 25, 50, 100}


def _paginate_log(db, table, search_cols, log_q, log_per_page_raw, log_page_raw, extra_where=None):
    """Filter and paginate any event-log table. Returns (entries, total, page, per_page, total_pages, page_list)."""
    if log_per_page_raw == 'all':
        log_per_page = 0
    else:
        try:
            log_per_page = int(log_per_page_raw)
        except (ValueError, TypeError):
            log_per_page = 25
        if log_per_page not in _VALID_LOG_PER_PAGE:
            log_per_page = 25
    try:
        log_page = max(1, int(log_page_raw or 1))
    except (ValueError, TypeError):
        log_page = 1

    log_params = []
    clauses = []
    if extra_where:
        clauses.append(extra_where)
    if log_q:
        clauses.append("(" + " OR ".join(f"{c} LIKE ?" for c in search_cols) + ")")
        log_params = [f'%{log_q}%'] * len(search_cols)
    log_where = ("WHERE " + " AND ".join(clauses)) if clauses else ""

    log_total = db.execute(
        f"SELECT COUNT(*) FROM {table} {log_where}", log_params
    ).fetchone()[0]

    if log_per_page:
        log_total_pages = max(1, (log_total + log_per_page - 1) // log_per_page)
        log_page = min(log_page, log_total_pages)
        entries = db.execute(
            f"SELECT * FROM {table} {log_where} ORDER BY timestamp DESC LIMIT ? OFFSET ?",
            log_params + [log_per_page, (log_page - 1) * log_per_page]
        ).fetchall()
        visible = sorted({1, log_total_pages} | set(range(max(1, log_page - 2), min(log_total_pages, log_page + 2) + 1)))
        log_page_list = []
        for i, p in enumerate(visible):
            if i > 0 and p > visible[i - 1] + 1:
                log_page_list.append(None)
            log_page_list.append(p)
    else:
        log_page, log_total_pages, log_page_list = 1, 1, [1]
        entries = db.execute(
            f"SELECT * FROM {table} {log_where} ORDER BY timestamp DESC", log_params
        ).fetchall()

    return entries, log_total, log_page, log_per_page, log_total_pages, log_page_list

@bp.route('/ref-docs')
def ref_docs_list():
    db = get_db()
    q          = request.args.get('q', '').strip()
    dist       = request.args.get('dist', '').strip()
    cui_filter = request.args.get('cui_filter', '').strip()
    expired    = request.args.get('expired', '').strip()
    sort       = request.args.get('sort', 'doc_number').strip()
    order      = request.args.get('order', 'asc').strip()
    per_page_raw = request.args.get('per_page', '25').strip()
    try:
        page = max(1, int(request.args.get('page', '1')))
    except ValueError:
        page = 1

    if sort not in _REF_DOC_SORT_COLS:
        sort = 'doc_number'
    if order not in ('asc', 'desc'):
        order = 'asc'
    if per_page_raw == 'all':
        per_page = 0
    else:
        try:
            per_page = int(per_page_raw)
        except ValueError:
            per_page = 25
        if per_page not in _VALID_REF_PER_PAGE:
            per_page = 25

    today = datetime.now().strftime('%Y-%m-%d')

    conditions, params = [], []
    if q:
        conditions.append("(doc_number LIKE ? OR title LIKE ? OR description LIKE ?)")
        params.extend([f'%{q}%', f'%{q}%', f'%{q}%'])
    if dist:
        conditions.append("dist_statement = ?")
        params.append(dist)
    if cui_filter == 'cui':
        conditions.append("cui = 1")
    elif cui_filter == 'non_cui':
        conditions.append("(cui = 0 OR cui IS NULL)")
    if expired == 'active':
        conditions.append("(expiration_date IS NULL OR expiration_date >= ?)")
        params.append(today)
    elif expired == 'expired':
        conditions.append("expiration_date IS NOT NULL AND expiration_date < ?")
        params.append(today)
    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""

    total = db.execute(
        f"SELECT COUNT(*) FROM reference_documents {where}", params
    ).fetchone()[0]

    if per_page:
        total_pages = max(1, (total + per_page - 1) // per_page)
        page = min(page, total_pages)
        offset = (page - 1) * per_page
        docs = db.execute(
            f"SELECT * FROM reference_documents {where}"
            f" ORDER BY {sort} {order.upper()} LIMIT ? OFFSET ?",
            params + [per_page, offset]
        ).fetchall()
        visible = sorted({1, total_pages} | set(range(max(1, page - 2), min(total_pages, page + 2) + 1)))
        page_list = []
        for i, p in enumerate(visible):
            if i > 0 and p > visible[i - 1] + 1:
                page_list.append(None)
            page_list.append(p)
    else:
        page, total_pages, page_list = 1, 1, [1]
        docs = db.execute(
            f"SELECT * FROM reference_documents {where} ORDER BY {sort} {order.upper()}",
            params
        ).fetchall()

    has_sensitive = db.execute(
        "SELECT COUNT(*) FROM reference_documents"
        " WHERE cui=1 OR (dist_statement IS NOT NULL AND dist_statement != 'A')"
    ).fetchone()[0] > 0

    cui_categories = db.execute(
        "SELECT code, label, description FROM cui_categories ORDER BY sort_order, code"
    ).fetchall()

    dissem_controls = db.execute(
        "SELECT code, label, description FROM cui_dissem_controls ORDER BY sort_order, code"
    ).fetchall()

    event_log, log_total, log_page, log_per_page, log_total_pages, log_page_list = _paginate_log(
        db, 'ref_doc_log',
        ['doc_number', 'action_type', 'description'],
        request.args.get('log_q', '').strip(),
        request.args.get('log_per_page', '25'),
        request.args.get('log_page', 1),
    )
    log_q = request.args.get('log_q', '').strip()

    return render_template(
        'ref_docs_config.html', docs=docs, today=today,
        has_sensitive=has_sensitive,
        cui_categories=cui_categories,
        dissem_controls=dissem_controls,
        dist_statements=DIST_STATEMENTS,
        event_log=event_log, log_q=log_q, log_per_page=log_per_page,
        log_page=log_page, log_total=log_total,
        log_total_pages=log_total_pages, log_page_list=log_page_list,
        sort=sort, order=order,
        q=q, dist=dist, cui_filter=cui_filter, expired=expired,
        per_page=per_page, page=page, total=total,
        total_pages=total_pages, page_list=page_list,
    )


@bp.route('/ref-docs/add', methods=['POST'])
def ref_doc_add():
    db = get_db()
    doc_number      = request.form.get('doc_number', '').strip()
    title           = request.form.get('title', '').strip()
    revision        = request.form.get('revision', '').strip()
    description     = request.form.get('description', '').strip()
    effective_date  = request.form.get('effective_date', '').strip()
    expiration_date = request.form.get('expiration_date', '').strip()
    url             = request.form.get('url', '').strip()
    cui             = 1 if request.form.get('cui') else 0
    cui_categories  = _parse_cui_categories(request.form)
    cui_dissem      = _parse_cui_dissem(request.form)
    dist_statement  = request.form.get('dist_statement', '').strip() or None
    dist_reason     = request.form.get('dist_reason', '').strip() or None
    if doc_number and title:
        row = db.execute("SELECT MAX(sort_order) as mx FROM reference_documents").fetchone()
        next_order = (row['mx'] or 0) + 1
        db.execute(
            "INSERT INTO reference_documents"
            " (doc_number, title, revision, description, effective_date, expiration_date,"
            "  url, cui, cui_categories, cui_dissem, dist_statement, dist_reason, sort_order)"
            " VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (doc_number, title, revision or None, description or None,
             effective_date or None, expiration_date or None, url or None,
             cui, cui_categories, cui_dissem, dist_statement, dist_reason, next_order)
        )
        doc_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]
        f = request.files.get('file')
        if f and f.filename:
            ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
            if ext in _ALLOWED_REF_DOC_EXT:
                stored = f"rdoc_{doc_id}_{uuid.uuid4().hex[:8]}.{ext}"
                f.save(os.path.join(_ref_doc_folder(), stored))
                db.execute(
                    "UPDATE reference_documents SET file_original=?, file_stored=? WHERE id=?",
                    (secure_filename(f.filename), stored, doc_id)
                )
        marking = _build_cui_marking(cui_categories, cui_dissem)
        parts = [title]
        if cui:
            parts.append(marking)
        if dist_statement:
            parts.append(f'Dist {dist_statement}')
        _log_ref_doc(db, doc_id, doc_number, 'Added', f"Added to library: {'; '.join(parts)}")
        db.commit()
        flash(f'"{doc_number}" added to reference library.', 'success')
    return redirect(url_for('ha.ref_docs_list'))


@bp.route('/ref-docs/<int:doc_id>/edit', methods=['POST'])
def ref_doc_edit(doc_id):
    db = get_db()
    doc_number      = request.form.get('doc_number', '').strip()
    title           = request.form.get('title', '').strip()
    revision        = request.form.get('revision', '').strip()
    description     = request.form.get('description', '').strip()
    effective_date  = request.form.get('effective_date', '').strip()
    expiration_date = request.form.get('expiration_date', '').strip()
    url             = request.form.get('url', '').strip()
    clear_file      = request.form.get('clear_file') == '1'
    cui             = 1 if request.form.get('cui') else 0
    cui_categories  = _parse_cui_categories(request.form)
    cui_dissem      = _parse_cui_dissem(request.form)
    dist_statement  = request.form.get('dist_statement', '').strip() or None
    dist_reason     = request.form.get('dist_reason', '').strip() or None
    if not (doc_number and title):
        return redirect(url_for('ha.ref_docs_list'))
    existing = db.execute("SELECT * FROM reference_documents WHERE id=?", (doc_id,)).fetchone()
    if not existing:
        return redirect(url_for('ha.ref_docs_list'))
    file_original = existing['file_original']
    file_stored   = existing['file_stored']
    f = request.files.get('file')
    if f and f.filename:
        ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
        if ext in _ALLOWED_REF_DOC_EXT:
            if file_stored:
                old_path = os.path.join(_ref_doc_folder(), file_stored)
                if os.path.exists(old_path):
                    os.remove(old_path)
            file_stored   = f"rdoc_{doc_id}_{uuid.uuid4().hex[:8]}.{ext}"
            file_original = secure_filename(f.filename)
            f.save(os.path.join(_ref_doc_folder(), file_stored))
    elif clear_file and file_stored:
        old_path = os.path.join(_ref_doc_folder(), file_stored)
        if os.path.exists(old_path):
            os.remove(old_path)
        file_original = None
        file_stored   = None
    db.execute(
        "UPDATE reference_documents"
        " SET doc_number=?, title=?, revision=?, description=?, effective_date=?, expiration_date=?,"
        "     url=?, file_original=?, file_stored=?,"
        "     cui=?, cui_categories=?, cui_dissem=?, dist_statement=?, dist_reason=?"
        " WHERE id=?",
        (doc_number, title, revision or None, description or None,
         effective_date or None, expiration_date or None, url or None,
         file_original, file_stored,
         cui, cui_categories, cui_dissem, dist_statement, dist_reason, doc_id)
    )
    changes = []
    if doc_number != existing['doc_number']:
        changes.append(f"Doc# → {doc_number}")
    if title != existing['title']:
        changes.append("Title updated")
    if (revision or '') != (existing['revision'] or ''):
        changes.append(f"Revision → {revision or '—'}")
    if (url or '') != (existing['url'] or ''):
        changes.append("URL updated")
    if bool(cui) != bool(existing['cui']):
        changes.append("CUI " + ("added" if cui else "removed"))
    elif cui and (cui_categories or '') != (existing['cui_categories'] or ''):
        changes.append("CUI categories updated")
    elif cui and (cui_dissem or '') != (existing.get('cui_dissem') or ''):
        changes.append("CUI dissemination controls updated")
    if (dist_statement or '') != (existing['dist_statement'] or ''):
        changes.append(f"Dist statement → {dist_statement or 'None'}")
    if file_stored and file_stored != existing['file_stored']:
        changes.append(f"File attached: {file_original}")
    elif clear_file:
        changes.append("File removed")
    desc = "Updated" + (f": {'; '.join(changes)}" if changes else " (no field changes)")
    _log_ref_doc(db, doc_id, doc_number, 'Edited', desc)
    db.commit()
    flash('Reference document updated.', 'success')
    return redirect(url_for('ha.ref_docs_list'))


@bp.route('/ref-docs/<int:doc_id>/delete', methods=['POST'])
def ref_doc_delete(doc_id):
    db = get_db()
    row = db.execute("SELECT * FROM reference_documents WHERE id = ?", (doc_id,)).fetchone()
    if not row:
        return redirect(url_for('ha.ref_docs_list'))
    in_use = db.execute(
        "SELECT COUNT(*) FROM hazard_references WHERE ref_doc_id = ?", (doc_id,)
    ).fetchone()[0]
    if in_use > 0:
        flash(
            f'Cannot delete "{row["doc_number"]}" — it is cited in {in_use} hazard analysis reference(s).',
            'error'
        )
        return redirect(url_for('ha.ref_docs_list'))
    _log_ref_doc(db, doc_id, row['doc_number'], 'Deleted', f"Removed from library: {row['title']}")
    if row['file_stored']:
        try:
            old_path = os.path.join(_ref_doc_folder(), row['file_stored'])
            if os.path.exists(old_path):
                os.remove(old_path)
        except OSError:
            pass
    db.execute("DELETE FROM reference_documents WHERE id = ?", (doc_id,))
    db.commit()
    flash(f'"{row["doc_number"]}" removed from reference library.', 'success')
    return redirect(url_for('ha.ref_docs_list'))


# ── CUI CATEGORY CONFIG ────────────────────────────────────────────────────────

_VALID_CUI_PER_PAGE  = {10, 25, 50, 100}
_VALID_CUI_SORT_COLS = {'code', 'label', 'sort_order'}


def _log_cui_marking(db, marking_type, code, action_type, description):
    db.execute(
        "INSERT INTO cui_marking_log (marking_type, code, timestamp, action_type, description)"
        " VALUES (?, ?, ?, ?, ?)",
        (marking_type, code, datetime.utcnow().isoformat(timespec='seconds'), action_type, description)
    )


def _cui_paginate(db, table, where, params, per_page, page, sort='sort_order', order='asc'):
    if sort not in _VALID_CUI_SORT_COLS:
        sort = 'sort_order'
    if order not in ('asc', 'desc'):
        order = 'asc'
    secondary = 'code' if sort != 'code' else 'sort_order'
    order_clause = f"{sort} {order.upper()}, {secondary}"
    total = db.execute(f"SELECT COUNT(*) FROM {table} {where}", params).fetchone()[0]
    if per_page:
        total_pages = max(1, (total + per_page - 1) // per_page)
        page = min(max(1, page), total_pages)
        offset = (page - 1) * per_page
        rows = db.execute(
            f"SELECT * FROM {table} {where} ORDER BY {order_clause} LIMIT ? OFFSET ?",
            params + [per_page, offset]
        ).fetchall()
        visible = sorted({1, total_pages} | set(range(max(1, page - 2), min(total_pages, page + 2) + 1)))
        page_list = []
        for i, p in enumerate(visible):
            if i > 0 and p > visible[i - 1] + 1:
                page_list.append(None)
            page_list.append(p)
    else:
        page, total_pages, page_list = 1, 1, [1]
        rows = db.execute(
            f"SELECT * FROM {table} {where} ORDER BY {order_clause}", params
        ).fetchall()
    return rows, total, page, total_pages, page_list


@bp.route('/cui-categories')
def cui_categories_list():
    db = get_db()
    q            = request.args.get('q', '').strip()
    per_page_raw = request.args.get('per_page', 'all')
    cat_page     = max(1, int(request.args.get('cat_page',    1) or 1))
    dissem_page  = max(1, int(request.args.get('dissem_page', 1) or 1))
    cat_sort     = request.args.get('cat_sort',    'sort_order')
    cat_order    = request.args.get('cat_order',   'asc')
    dissem_sort  = request.args.get('dissem_sort', 'sort_order')
    dissem_order = request.args.get('dissem_order','asc')

    if per_page_raw == 'all':
        per_page = 0
    else:
        try:
            per_page = int(per_page_raw)
        except ValueError:
            per_page = 25
        if per_page not in _VALID_CUI_PER_PAGE:
            per_page = 25

    if q:
        where  = "WHERE (code LIKE ? OR label LIKE ? OR description LIKE ?)"
        params = [f'%{q}%', f'%{q}%', f'%{q}%']
    else:
        where, params = '', []

    cats,   cat_total,   cat_page,   cat_total_pages,   cat_page_list   = _cui_paginate(
        db, 'cui_categories',     where, params[:], per_page, cat_page,   cat_sort,    cat_order)
    dissems, dissem_total, dissem_page, dissem_total_pages, dissem_page_list = _cui_paginate(
        db, 'cui_dissem_controls', where, params[:], per_page, dissem_page, dissem_sort, dissem_order)

    event_log, log_total, log_page, log_per_page, log_total_pages, log_page_list = _paginate_log(
        db, 'cui_marking_log',
        ['code', 'marking_type', 'action_type', 'description'],
        request.args.get('log_q', '').strip(),
        request.args.get('log_per_page', '25'),
        request.args.get('log_page', 1),
    )
    log_q = request.args.get('log_q', '').strip()

    return render_template(
        'cui_categories_config.html',
        cats=cats, cat_total=cat_total, cat_page=cat_page,
        cat_total_pages=cat_total_pages, cat_page_list=cat_page_list,
        cat_sort=cat_sort, cat_order=cat_order,
        dissems=dissems, dissem_total=dissem_total, dissem_page=dissem_page,
        dissem_total_pages=dissem_total_pages, dissem_page_list=dissem_page_list,
        dissem_sort=dissem_sort, dissem_order=dissem_order,
        q=q, per_page=per_page,
        event_log=event_log, log_q=log_q, log_per_page=log_per_page,
        log_page=log_page, log_total=log_total,
        log_total_pages=log_total_pages, log_page_list=log_page_list,
    )


@bp.route('/cui-categories/add', methods=['POST'])
def cui_category_add():
    db    = get_db()
    code  = request.form.get('code', '').strip().upper()
    label = request.form.get('label', '').strip()
    if code and label:
        description    = request.form.get('description', '').strip() or None
        sort_order_str = request.form.get('sort_order', '').strip()
        if sort_order_str.isdigit():
            sort_order = int(sort_order_str)
        else:
            row = db.execute("SELECT MAX(sort_order) as mx FROM cui_categories").fetchone()
            sort_order = (row['mx'] or 0) + 1
        db.execute(
            "INSERT OR IGNORE INTO cui_categories (code, label, description, sort_order) VALUES (?, ?, ?, ?)",
            (code, label, description, sort_order)
        )
        _log_cui_marking(db, 'category', code, 'Added', f'Added category: {label}')
        db.commit()
        flash(f'Category {code} added.', 'success')
    return redirect(url_for('ha.cui_categories_list'))


@bp.route('/cui-categories/<int:cat_id>/edit', methods=['POST'])
def cui_category_edit(cat_id):
    db    = get_db()
    code  = request.form.get('code', '').strip().upper()
    label = request.form.get('label', '').strip()
    if code and label:
        description    = request.form.get('description', '').strip() or None
        sort_order_str = request.form.get('sort_order', '').strip()
        if sort_order_str.isdigit():
            db.execute(
                "UPDATE cui_categories SET code=?, label=?, description=?, sort_order=? WHERE id=?",
                (code, label, description, int(sort_order_str), cat_id)
            )
        else:
            db.execute(
                "UPDATE cui_categories SET code=?, label=?, description=? WHERE id=?",
                (code, label, description, cat_id)
            )
        _log_cui_marking(db, 'category', code, 'Edited', f'Updated category: {label}')
        db.commit()
        flash('Category updated.', 'success')
    return redirect(url_for('ha.cui_categories_list'))


@bp.route('/cui-categories/<int:cat_id>/delete', methods=['POST'])
def cui_category_delete(cat_id):
    db  = get_db()
    row = db.execute("SELECT * FROM cui_categories WHERE id=?", (cat_id,)).fetchone()
    if not row:
        return redirect(url_for('ha.cui_categories_list'))
    in_use = db.execute(
        "SELECT COUNT(*) FROM reference_documents"
        " WHERE ',' || cui_categories || ',' LIKE ?",
        (f'%,{row["code"]},%',)
    ).fetchone()[0]
    if in_use:
        flash(f'{row["code"]} is cited in {in_use} document(s) and cannot be removed.', 'error')
    else:
        _log_cui_marking(db, 'category', row['code'], 'Removed', f'Removed category: {row["label"]}')
        db.execute("DELETE FROM cui_categories WHERE id=?", (cat_id,))
        db.commit()
        flash(f'{row["code"]} removed.', 'success')
    return redirect(url_for('ha.cui_categories_list'))


# ── CUI DISSEMINATION CONTROL CONFIG ──────────────────────────────────────────

@bp.route('/cui-dissem/add', methods=['POST'])
def cui_dissem_add():
    db    = get_db()
    code  = request.form.get('code', '').strip().upper()
    label = request.form.get('label', '').strip()
    if code and label:
        description    = request.form.get('description', '').strip() or None
        sort_order_str = request.form.get('sort_order', '').strip()
        if sort_order_str.isdigit():
            sort_order = int(sort_order_str)
        else:
            row = db.execute("SELECT MAX(sort_order) as mx FROM cui_dissem_controls").fetchone()
            sort_order = (row['mx'] or 0) + 1
        db.execute(
            "INSERT OR IGNORE INTO cui_dissem_controls (code, label, description, sort_order) VALUES (?, ?, ?, ?)",
            (code, label, description, sort_order)
        )
        _log_cui_marking(db, 'dissemination', code, 'Added', f'Added dissemination control: {label}')
        db.commit()
        flash(f'Dissemination control {code} added.', 'success')
    return redirect(url_for('ha.cui_categories_list'))


@bp.route('/cui-dissem/<int:ctrl_id>/edit', methods=['POST'])
def cui_dissem_edit(ctrl_id):
    db    = get_db()
    code  = request.form.get('code', '').strip().upper()
    label = request.form.get('label', '').strip()
    if code and label:
        description    = request.form.get('description', '').strip() or None
        sort_order_str = request.form.get('sort_order', '').strip()
        if sort_order_str.isdigit():
            db.execute(
                "UPDATE cui_dissem_controls SET code=?, label=?, description=?, sort_order=? WHERE id=?",
                (code, label, description, int(sort_order_str), ctrl_id)
            )
        else:
            db.execute(
                "UPDATE cui_dissem_controls SET code=?, label=?, description=? WHERE id=?",
                (code, label, description, ctrl_id)
            )
        _log_cui_marking(db, 'dissemination', code, 'Edited', f'Updated dissemination control: {label}')
        db.commit()
        flash('Dissemination control updated.', 'success')
    return redirect(url_for('ha.cui_categories_list'))


@bp.route('/cui-dissem/<int:ctrl_id>/delete', methods=['POST'])
def cui_dissem_delete(ctrl_id):
    db  = get_db()
    row = db.execute("SELECT * FROM cui_dissem_controls WHERE id=?", (ctrl_id,)).fetchone()
    if not row:
        return redirect(url_for('ha.cui_categories_list'))
    in_use = db.execute(
        "SELECT COUNT(*) FROM reference_documents"
        " WHERE ',' || cui_dissem || ',' LIKE ?",
        (f'%,{row["code"]},%',)
    ).fetchone()[0]
    if in_use:
        flash(f'{row["code"]} is applied to {in_use} document(s) and cannot be removed.', 'error')
    else:
        _log_cui_marking(db, 'dissemination', row['code'], 'Removed', f'Removed dissemination control: {row["label"]}')
        db.execute("DELETE FROM cui_dissem_controls WHERE id=?", (ctrl_id,))
        db.commit()
        flash(f'{row["code"]} removed.', 'success')
    return redirect(url_for('ha.cui_categories_list'))


# ── PER-HA REFERENCES ─────────────────────────────────────────────────────────

@bp.route('/<int:ha_pk>/refs/add', methods=['POST'])
def ha_ref_add(ha_pk):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    signatures = db.execute("SELECT * FROM hazard_signatures WHERE ha_id = ?", (ha_pk,)).fetchall()
    sig_map = {s['role']: dict(s) for s in signatures}
    if refs_locked(ha, sig_map):
        flash('This hazard analysis is locked — references cannot be modified.', 'error')
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    ref_doc_id    = request.form.get('ref_doc_id', '').strip()
    section_cited = request.form.get('section_cited', '').strip()
    if not ref_doc_id:
        flash('Please select a document from the library.', 'error')
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    lib_doc = db.execute(
        "SELECT * FROM reference_documents WHERE id = ?", (ref_doc_id,)
    ).fetchone()
    if not lib_doc:
        flash('Document not found in library.', 'error')
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    row = db.execute(
        "SELECT MAX(sort_order) as mx FROM hazard_references WHERE ha_id = ?", (ha_pk,)
    ).fetchone()
    next_order = (row['mx'] or 0) + 1
    db.execute(
        "INSERT INTO hazard_references (ha_id, ref_doc_id, doc_number, title, revision, section_cited, sort_order)"
        " VALUES (?,?,?,?,?,?,?)",
        (ha_pk, lib_doc['id'], lib_doc['doc_number'], lib_doc['title'],
         lib_doc['revision'], section_cited or None, next_order)
    )
    _log_ha(db, ha_pk, ha['ha_id'], 'Reference Added',
            f"Added reference: {lib_doc['doc_number']} — {lib_doc['title']}")
    db.commit()
    flash('Reference added.', 'success')
    return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')


@bp.route('/<int:ha_pk>/refs/add-new', methods=['POST'])
def ha_ref_add_new(ha_pk):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    signatures = db.execute("SELECT * FROM hazard_signatures WHERE ha_id = ?", (ha_pk,)).fetchall()
    sig_map = {s['role']: dict(s) for s in signatures}
    if refs_locked(ha, sig_map):
        flash('This hazard analysis is locked — references cannot be modified.', 'error')
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    doc_number    = request.form.get('doc_number', '').strip()
    title         = request.form.get('title', '').strip()
    revision      = request.form.get('revision', '').strip() or None
    section_cited = request.form.get('section_cited', '').strip() or None
    dist_statement = request.form.get('dist_statement', '').strip() or None
    dist_reason    = request.form.get('dist_reason', '').strip() or None
    cui            = 1 if request.form.get('cui') else 0
    cui_cats       = _parse_cui_categories(request.form) if cui else None
    cui_dissem     = _parse_cui_dissem(request.form) if cui else None
    if not doc_number or not title:
        flash('Document number and title are required.', 'error')
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    row = db.execute("SELECT MAX(sort_order) as mx FROM reference_documents").fetchone()
    db.execute(
        "INSERT INTO reference_documents"
        " (doc_number, title, revision, cui, cui_categories, cui_dissem,"
        "  dist_statement, dist_reason, sort_order)"
        " VALUES (?,?,?,?,?,?,?,?,?)",
        (doc_number, title, revision, cui, cui_cats, cui_dissem,
         dist_statement, dist_reason, (row['mx'] or 0) + 1)
    )
    lib_doc_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]
    marking = _build_cui_marking(cui_cats, cui_dissem)
    desc_parts = [title]
    if cui:       desc_parts.append(marking)
    if dist_statement: desc_parts.append(f'Dist {dist_statement}')
    _log_ref_doc(db, lib_doc_id, doc_number, 'Added',
                 f"Added from HA {ha['ha_id']}: {'; '.join(desc_parts)}")
    row = db.execute("SELECT MAX(sort_order) as mx FROM hazard_references WHERE ha_id = ?", (ha_pk,)).fetchone()
    db.execute(
        "INSERT INTO hazard_references"
        " (ha_id, ref_doc_id, doc_number, title, revision, section_cited, sort_order)"
        " VALUES (?,?,?,?,?,?,?)",
        (ha_pk, lib_doc_id, doc_number, title, revision, section_cited, (row['mx'] or 0) + 1)
    )
    _log_ha(db, ha_pk, ha['ha_id'], 'Reference Added',
            f"Added new doc to library and referenced: {doc_number} — {title}")
    db.commit()
    flash(f'"{doc_number}" added to library and referenced.', 'success')
    return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')


@bp.route('/<int:ha_pk>/refs/<int:ref_id>/delete', methods=['POST'])
def ha_ref_delete(ha_pk, ref_id):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    signatures = db.execute("SELECT * FROM hazard_signatures WHERE ha_id = ?", (ha_pk,)).fetchall()
    sig_map = {s['role']: dict(s) for s in signatures}
    if refs_locked(ha, sig_map):
        flash('This hazard analysis is locked — references cannot be modified.', 'error')
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    ref = db.execute("SELECT doc_number, title FROM hazard_references WHERE id=? AND ha_id=?",
                     (ref_id, ha_pk)).fetchone()
    db.execute("DELETE FROM hazard_references WHERE id = ? AND ha_id = ?", (ref_id, ha_pk))
    if ref:
        _log_ha(db, ha_pk, ha['ha_id'], 'Reference Removed',
                f"Removed reference: {ref['doc_number']} — {ref['title']}")
    db.commit()
    return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')


@bp.route('/<int:ha_pk>/refs/<int:ref_id>/move', methods=['POST'])
def ha_ref_move(ha_pk, ref_id):
    db = get_db()
    ha = get_ha_or_404(db, ha_pk)
    signatures = db.execute("SELECT * FROM hazard_signatures WHERE ha_id = ?", (ha_pk,)).fetchall()
    sig_map = {s['role']: dict(s) for s in signatures}
    if refs_locked(ha, sig_map):
        flash('This hazard analysis is locked — references cannot be modified.', 'error')
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    direction = request.form.get('direction')
    refs = db.execute(
        "SELECT id FROM hazard_references WHERE ha_id = ? ORDER BY sort_order, id", (ha_pk,)
    ).fetchall()
    ids = [r['id'] for r in refs]
    try:
        idx = ids.index(ref_id)
    except ValueError:
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    if direction == 'up' and idx > 0:
        ids[idx], ids[idx - 1] = ids[idx - 1], ids[idx]
    elif direction == 'down' and idx < len(ids) - 1:
        ids[idx], ids[idx + 1] = ids[idx + 1], ids[idx]
    else:
        return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
    for order, rid in enumerate(ids, 1):
        db.execute("UPDATE hazard_references SET sort_order=? WHERE id=?", (order, rid))
    db.commit()
    return redirect(url_for('ha.ha_detail', ha_pk=ha_pk) + '#references')
